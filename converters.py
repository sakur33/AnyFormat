"""
Lógica de conversión por categorías de formato.
Cada conversor recibe (ruta_origen, ruta_destino) y produce el archivo destino.
"""
import os

# Registrar apertura de HEIC/HEIF en Pillow (formato de fotos de iPhone).
try:
    import pillow_heif
    pillow_heif.register_heif_opener()
except Exception:
    pass

# ---------------------------------------------------------------------------
# Definición de formatos soportados por categoría
# ---------------------------------------------------------------------------
# Formatos raster que Pillow puede leer Y escribir.
IMAGE_FORMATS = [
    "png", "jpg", "jpeg", "bmp", "gif", "tiff", "webp", "ico",
    "heic", "heif", "avif",
]
# SVG es vectorial: solo origen (se rasteriza). No es destino válido.
VECTOR_FORMATS = ["svg"]
DOC_FORMATS = ["pdf", "docx", "txt", "html", "md"]
AUDIO_FORMATS = ["mp3", "wav", "ogg", "flac", "aac", "m4a"]
VIDEO_FORMATS = ["mp4", "avi", "mkv", "mov", "webm", "gif"]
DATA_FORMATS = ["csv", "xlsx", "json", "tsv"]

ALL_FORMATS = sorted(set(
    IMAGE_FORMATS + VECTOR_FORMATS + DOC_FORMATS
    + AUDIO_FORMATS + VIDEO_FORMATS + DATA_FORMATS
))

# Orden de presentación en los desplegables. Un formato aparece una sola vez,
# en la primera categoría que lo reclame (gif es imagen y vídeo: gana imagen).
CATEGORIES = [
    ("Imagen", IMAGE_FORMATS),
    ("Vectorial", VECTOR_FORMATS),
    ("Documento", DOC_FORMATS),
    ("Audio", AUDIO_FORMATS),
    ("Vídeo", VIDEO_FORMATS),
    ("Datos", DATA_FORMATS),
]


def _category(fmt):
    fmt = fmt.lower().lstrip(".")
    if fmt in IMAGE_FORMATS:
        return "image"
    if fmt in VECTOR_FORMATS:
        return "vector"
    if fmt in DOC_FORMATS:
        return "doc"
    if fmt in AUDIO_FORMATS:
        return "audio"
    if fmt in VIDEO_FORMATS:
        return "video"
    if fmt in DATA_FORMATS:
        return "data"
    return None


def can_convert(src_fmt, dst_fmt):
    """Devuelve True si la conversión entre dos formatos está soportada."""
    src_fmt = src_fmt.lower().lstrip(".")
    dst_fmt = dst_fmt.lower().lstrip(".")
    if src_fmt == dst_fmt:
        return False
    c_src, c_dst = _category(src_fmt), _category(dst_fmt)
    if c_src is None or c_dst is None:
        return False
    # Conversiones dentro de la misma categoría
    if c_src == c_dst:
        return True
    # imagen/vector -> doc: solo a PDF (única ruta implementada)
    if c_src in ("image", "vector") and c_dst == "doc":
        return dst_fmt == "pdf"
    # doc -> imagen: solo desde PDF
    if c_src == "doc" and c_dst == "image":
        return src_fmt == "pdf"
    # Cruces permitidos entre categorías
    cross = {
        ("video", "image"),  # video -> gif/frame
        ("video", "audio"),  # extraer audio
        ("vector", "image"), # svg -> png/jpg (rasterizado)
    }
    return (c_src, c_dst) in cross


def valid_targets(src_fmt):
    """Lista de formatos destino válidos para un formato origen dado."""
    return [f for f in ALL_FORMATS if can_convert(src_fmt, f)]


def _group(formats):
    """Agrupa formatos por categoría, sin repetir ninguno entre grupos.
    Devuelve [(categoría, [formatos ordenados]), ...] omitiendo grupos vacíos."""
    allowed, seen, groups = set(formats), set(), []
    for label, fmts in CATEGORIES:
        picked = sorted(f for f in fmts if f in allowed and f not in seen)
        if picked:
            seen.update(picked)
            groups.append((label, picked))
    return groups


def all_formats_by_category():
    """Todos los formatos agrupados, para el desplegable de origen."""
    return _group(ALL_FORMATS)


def targets_by_category(src_fmt):
    """Formatos destino válidos agrupados, para el desplegable de destino."""
    return _group(valid_targets(src_fmt))


def detect_format(path):
    """Extensión normalizada de un archivo, o "" si no está soportada."""
    ext = os.path.splitext(path)[1].lstrip(".").lower()
    return ext if ext in ALL_FORMATS else ""


# ---------------------------------------------------------------------------
# Conversores concretos
# ---------------------------------------------------------------------------
def convert_image(src, dst, dst_fmt):
    from PIL import Image
    img = Image.open(src)
    save_fmt = dst_fmt.upper()
    if save_fmt == "JPG":
        save_fmt = "JPEG"
    if save_fmt in ("HEIC", "HEIF"):
        save_fmt = "HEIF"  # el plugin pillow-heif registra el escritor como HEIF
    if save_fmt in ("JPEG", "BMP", "HEIF") and img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")
    if save_fmt == "ICO":
        img.save(dst, format="ICO", sizes=[(256, 256)])
    else:
        img.save(dst, format=save_fmt)


def convert_image_to_pdf(src, dst):
    from PIL import Image
    img = Image.open(src)
    if img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")
    img.save(dst, "PDF", resolution=100.0)


def convert_pdf_to_image(src, dst, dst_fmt):
    # Renderiza la primera página como imagen usando PyMuPDF
    import fitz
    doc = fitz.open(src)
    page = doc.load_page(0)
    pix = page.get_pixmap(dpi=150)
    pix.save(dst)
    doc.close()


def convert_doc(src, dst, src_fmt, dst_fmt):
    """Conversiones de documentos basadas en texto."""
    src_fmt, dst_fmt = src_fmt.lower(), dst_fmt.lower()

    def read_text():
        if src_fmt == "txt" or src_fmt == "md":
            with open(src, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        if src_fmt == "docx":
            from docx import Document
            return "\n".join(p.text for p in Document(src).paragraphs)
        if src_fmt == "html":
            from bs4 import BeautifulSoup
            with open(src, "r", encoding="utf-8", errors="replace") as f:
                return BeautifulSoup(f.read(), "html.parser").get_text("\n")
        if src_fmt == "pdf":
            import fitz
            doc = fitz.open(src)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text
        raise ValueError(f"Lectura no soportada: {src_fmt}")

    text = read_text()

    if dst_fmt in ("txt", "md"):
        with open(dst, "w", encoding="utf-8") as f:
            f.write(text)
    elif dst_fmt == "html":
        body = "".join(f"<p>{line}</p>\n" for line in text.splitlines())
        with open(dst, "w", encoding="utf-8") as f:
            f.write(f"<!DOCTYPE html>\n<html><body>\n{body}</body></html>")
    elif dst_fmt == "docx":
        from docx import Document
        d = Document()
        for line in text.splitlines():
            d.add_paragraph(line)
        d.save(dst)
    elif dst_fmt == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        c = canvas.Canvas(dst, pagesize=A4)
        width, height = A4
        y = height - 2 * cm
        for line in text.splitlines():
            if y < 2 * cm:
                c.showPage()
                y = height - 2 * cm
            c.drawString(2 * cm, y, line[:110])
            y -= 0.6 * cm
        c.save()
    else:
        raise ValueError(f"Destino no soportado: {dst_fmt}")


def convert_media(src, dst, dst_fmt):
    """Audio y video mediante ffmpeg (imageio-ffmpeg incluye el binario)."""
    import subprocess
    import imageio_ffmpeg
    ffmpeg = imageio_ffmpeg.get_ffmpeg_exe()
    cmd = [ffmpeg, "-y", "-i", src]
    if dst_fmt in AUDIO_FORMATS:
        cmd += ["-vn"]  # solo audio
    cmd.append(dst)
    creationflags = 0x08000000 if os.name == "nt" else 0  # sin ventana de consola
    proc = subprocess.run(
        cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        creationflags=creationflags,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.decode("utf-8", errors="replace")[-500:])


def convert_data(src, dst, src_fmt, dst_fmt):
    """Datos tabulares con pandas."""
    import pandas as pd
    src_fmt, dst_fmt = src_fmt.lower(), dst_fmt.lower()

    if src_fmt == "csv":
        df = pd.read_csv(src)
    elif src_fmt == "tsv":
        df = pd.read_csv(src, sep="\t")
    elif src_fmt == "xlsx":
        df = pd.read_excel(src)
    elif src_fmt == "json":
        df = pd.read_json(src)
    else:
        raise ValueError(f"Lectura no soportada: {src_fmt}")

    if dst_fmt == "csv":
        df.to_csv(dst, index=False)
    elif dst_fmt == "tsv":
        df.to_csv(dst, sep="\t", index=False)
    elif dst_fmt == "xlsx":
        df.to_excel(dst, index=False)
    elif dst_fmt == "json":
        df.to_json(dst, orient="records", force_ascii=False, indent=2)
    else:
        raise ValueError(f"Destino no soportado: {dst_fmt}")


def convert_svg(src, dst, dst_fmt):
    """Rasteriza SVG a imagen o PDF mediante cairosvg."""
    import cairosvg
    dst_fmt = dst_fmt.lower()
    if dst_fmt == "png":
        cairosvg.svg2png(url=src, write_to=dst)
    elif dst_fmt == "pdf":
        cairosvg.svg2pdf(url=src, write_to=dst)
    else:
        # Para jpg/bmp/etc: rasterizar a PNG en memoria y reconvertir.
        import io
        from PIL import Image
        png_bytes = cairosvg.svg2png(url=src)
        img = Image.open(io.BytesIO(png_bytes))
        save_fmt = dst_fmt.upper()
        if save_fmt == "JPG":
            save_fmt = "JPEG"
        if save_fmt in ("JPEG", "BMP") and img.mode in ("RGBA", "P", "LA"):
            img = img.convert("RGB")
        if save_fmt == "ICO":
            img.save(dst, format="ICO", sizes=[(256, 256)])
        else:
            img.save(dst, format=save_fmt)


# ---------------------------------------------------------------------------
# Punto de entrada único
# ---------------------------------------------------------------------------
def convert_file(src, dst, src_fmt, dst_fmt):
    """Despacha a la función adecuada según las categorías implicadas."""
    src_fmt = src_fmt.lower().lstrip(".")
    dst_fmt = dst_fmt.lower().lstrip(".")
    c_src, c_dst = _category(src_fmt), _category(dst_fmt)

    if c_src == "image" and c_dst == "image":
        convert_image(src, dst, dst_fmt)
    elif c_src == "vector" and c_dst in ("image", "doc"):
        convert_svg(src, dst, dst_fmt)
    elif c_src == "image" and c_dst == "doc":
        convert_image_to_pdf(src, dst)
    elif c_src == "doc" and c_dst == "image":
        convert_pdf_to_image(src, dst, dst_fmt)
    elif c_src == "doc" and c_dst == "doc":
        convert_doc(src, dst, src_fmt, dst_fmt)
    elif c_src in ("audio", "video") and c_dst in ("audio", "video", "image"):
        convert_media(src, dst, dst_fmt)
    elif c_src == "data" and c_dst == "data":
        convert_data(src, dst, src_fmt, dst_fmt)
    else:
        raise ValueError(f"Conversión {src_fmt} -> {dst_fmt} no soportada")
